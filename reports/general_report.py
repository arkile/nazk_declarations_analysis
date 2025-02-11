import os
from dataclasses import dataclass
from enum import Enum
from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, shared
from docx.shared import RGBColor


class ReportLevel(Enum):
    TOP = 1
    STEP = 2
    SUBSTEP = 3
    DETAILS = 4


@dataclass
class Entry:
    text: str
    level: ReportLevel
    criticality: int
    hyperlink: str = None


class GeneralReport(object):

    def __init__(self):
        self.full_name = None
        self.period = None
        self.entry_list: list[Entry] = []
        # self.text = ''
        # self.summary = ''


    def add_header(self, full_name: str, period: tuple[int, int]):
        self.full_name = full_name
        # self.text += '\n' + f'{period[0]} - {period[1]}'
        self.period = period

    # TODO for backward compatibility, remove later
    def add_top_info(self, full_name: str, period: tuple[int, int]):
        self.add_header(full_name, period)

    # critical level:
    # 1 - regular info
    # 2 - questionable information - may be a risk, may not
    # 3 - very likely to be a risk
    def add_record(self, report_level: ReportLevel, line: str, critical: int = 1, hyperlink = None):
        self.entry_list.append(Entry(text=line, level=report_level, criticality=critical, hyperlink=hyperlink))


    def print_to_docx(self):
        doc = Document('template.docx')
        # style = doc.styles['Normal']
        # font = style.font
        # font.name = 'Arial'
        title = doc.add_heading(self.full_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        years_heading = doc.add_heading(f'{self.period[0]} - {self.period[1]}', level=2)
        years_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # doc.add_paragraph(self.full_name)

        for entry in self.entry_list:
            match entry.level:
                case ReportLevel.TOP:
                    # step = doc.add_heading(entry.text, level=3)
                    step = doc.add_heading(level=3)
                    run = step.add_run(entry.text)
                    run.bold = True
                    if entry.hyperlink is not None:
                        # run.add_hyperlink(entry.hyperlink)
                        add_hyperlink_into_run(step, run, entry.hyperlink)
                case ReportLevel.STEP:
                    paragraph = doc.add_heading(level=5)
                    paragraph.add_run(entry.text).bold = True
                case ReportLevel.SUBSTEP:
                    paragraph = doc.add_heading(level=6)
                    run = paragraph.add_run(entry.text)
                    if entry.criticality == 3:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                case ReportLevel.DETAILS:
                    paragraph = doc.add_heading(level=7)
                    run = paragraph.add_run(entry.text)
                    if entry.criticality == 3:
                        run.font.color.rgb = RGBColor(255, 0, 0)
        if os.path.exists(f'doc_reports/{self.full_name}.docx'):
            os.remove(f'doc_reports/{self.full_name}.docx')
        os.makedirs('doc_reports', exist_ok=True)
        doc.save(f'doc_reports/{self.full_name}.docx')


    def __str__(self):
        text = self.full_name
        text += '\n' + f'{self.period[0]} - {self.period[1]}'
        for entry_ in self.entry_list:
            text += '\n'
            text += '\t' * (entry_.level.value - 1)
            line_ = entry_.text.replace('\n', '\n' + ('\t' * (entry_.level.value - 1)), -1)
            if entry_.criticality == 2:
                line_ = '! ' + line_ + ' !'
                text += line_
                # self.summary += f'\n {line_}'
            elif entry_.criticality == 3:
                line_ = line_ + ' !!!!!!!'
                text += line_
                # self.summary += f'\n {line_}'
            else:
                text += line_
        return text


    def __repr__(self):
        #TODO rewrite
        return self.__str__()
    # --- class Report end ---


def init_new_report():
    return GeneralReport()


def add_hyperlink_into_run(paragraph, run, url):
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    paragraph._p.insert(i+1,hyperlink)